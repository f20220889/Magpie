"""Extract a candidate's skills from an uploaded resume.

Supports PDF, Word (.docx) and plain text via pure-Python parsers, and images
(PNG/JPG/…) via a vision-capable LLM (OCR through the model). The parsed text
or image is handed to the configured LLM, which returns skills grouped under a
few broad domains — the same ``{domains: [{name, skills}]}`` shape the profile
uses, so the verified result can be saved straight through ``/api/init``.

Nothing here touches existing flows; it is an additive feature module.
"""

from __future__ import annotations

import base64
import io

from magpie.llm.base import LLMClient
from magpie.llm.factory import get_llm

# Cap text fed to the model so a long resume can't blow the prompt budget.
_MAX_CHARS = 16000

# Image types we can hand to a vision model.
_IMAGE_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
    "bmp": "image/bmp",
}

_SYSTEM = (
    "You extract a tech professional's SKILLS from their resume. Group the "
    "skills under a few broad domains (e.g. Backend, Frontend, DevOps, "
    "Data/ML, Cloud, Mobile, Security, QA). Include programming languages, "
    "frameworks, libraries, tools, platforms, databases and methodologies that "
    "are actually present. Do NOT invent skills the resume does not support. "
    "Use short, canonical names (e.g. 'PostgreSQL', not 'Postgres database "
    "experience'). Respond ONLY with JSON of the form: "
    '{"domains": [{"name": "...", "skills": ["...", "..."]}]}'
)


class IngestError(Exception):
    """Raised when an upload can't be read (unsupported type, empty, corrupt)."""


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def is_image(filename: str, content_type: str = "") -> bool:
    if content_type.lower().startswith("image/"):
        return True
    return _ext(filename) in _IMAGE_MIME


def _image_mime(filename: str, content_type: str) -> str:
    if content_type.lower().startswith("image/"):
        return content_type.lower()
    return _IMAGE_MIME.get(_ext(filename), "image/png")


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:  # encrypted / corrupt / not a real PDF
        raise IngestError(f"could not read PDF: {e}") from e


def _docx_text(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise IngestError(f"could not read Word document: {e}") from e
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:  # skills often live in tables
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    """Extract readable text from a non-image resume (PDF / DOCX / plain text)."""
    ext = _ext(filename)
    if ext == "pdf" or content_type == "application/pdf":
        return _pdf_text(data)
    if ext == "docx" or "wordprocessingml" in content_type:
        return _docx_text(data)
    if ext in ("txt", "md", "text") or content_type.startswith("text/"):
        return data.decode("utf-8", "ignore")
    raise IngestError(
        f"unsupported file type '{ext or content_type or 'unknown'}'. "
        "Upload a PDF, Word (.docx), image, or text resume."
    )


def normalize_skills(data: object) -> dict:
    """Coerce raw LLM output into ``{domains: [{name, skills:[...]}]}``.

    Tolerant of shape drift: drops non-strings, empty names/skill-lists, and
    de-dupes case-insensitively while preserving the model's ordering. A flat
    ``{skills: [...]}`` reply is bucketed under a single 'General' domain.
    """
    domains_raw: list = []
    if isinstance(data, dict):
        if isinstance(data.get("domains"), list):
            domains_raw = data["domains"]
        elif isinstance(data.get("skills"), list):
            domains_raw = [{"name": "General", "skills": data["skills"]}]

    out: list[dict] = []
    seen_domains: set[str] = set()
    for d in domains_raw:
        if not isinstance(d, dict):
            continue
        name = str(d.get("name", "")).strip()
        if not name or name.lower() in seen_domains:
            continue
        skills: list[str] = []
        seen_skills: set[str] = set()
        raw_skills = d.get("skills", [])
        if isinstance(raw_skills, list):
            for s in raw_skills:
                s = str(s).strip()
                if s and s.lower() not in seen_skills:
                    seen_skills.add(s.lower())
                    skills.append(s)
        if not skills:
            continue
        seen_domains.add(name.lower())
        out.append({"name": name, "skills": skills})
    return {"domains": out}


class SkillExtractor:
    def __init__(self, llm: LLMClient | None = None) -> None:
        self.llm = llm or get_llm()

    def from_text(self, text: str) -> dict:
        text = (text or "").strip()
        if not text:
            raise IngestError("no readable text found in the file")
        user = f'Resume text:\n"""\n{text[:_MAX_CHARS]}\n"""\n\nExtract the skills as JSON.'
        return normalize_skills(self.llm.complete_json(user, system=_SYSTEM))

    def from_image(self, image_b64: str, mime: str) -> dict:
        user = (
            "This image is a resume. Read it and extract the candidate's "
            "technical skills as JSON."
        )
        return normalize_skills(
            self.llm.complete_vision_json(user, image_b64, mime, system=_SYSTEM)
        )


def extract_skills_from_upload(
    filename: str, content_type: str, data: bytes, llm: LLMClient | None = None
) -> dict:
    """Dispatch an uploaded resume to the text or vision path -> normalized skills."""
    if not data:
        raise IngestError("the uploaded file is empty")
    extractor = SkillExtractor(llm=llm)
    if is_image(filename, content_type):
        b64 = base64.b64encode(data).decode("ascii")
        return extractor.from_image(b64, _image_mime(filename, content_type))
    return extractor.from_text(extract_text(filename, content_type, data))
